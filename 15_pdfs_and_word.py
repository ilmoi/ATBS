# note even though the package is called python-docx, when importing it we just say docs
import docx
import PyPDF2

pdfFileObj = open('automate_online-materials/meetingminutes.pdf', 'rb')  # read binary mode
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

# get number of pages
print(pdfReader.numPages)

# get a page
pageObj = pdfReader.getPage(0)  # note the first page is 0
print(pageObj.extractText())

pdfFileObj.close()


print("-------------------- ENCRYPTED --------------------")
pdfFileObj = open('automate_online-materials/encrypted.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
print(pdfReader.isEncrypted)  # true

# pdfReader.getPage(0)  # gives an error because file encrypted
pdfReader.decrypt('rosebud')  # status 1 = success
pageObj = pdfReader.getPage(0)
print(pageObj.extractText())

pdfFileObj.close()


print("-------------------- WRITER --------------------")
# can't write arbitrary text like with other types of files - can only copy/rotate/overlay pages from other PDFs!
# in fact you can't even edit the existing pdf - you have to copy the contents of the existing pdf into a new one

# open both files
pdf1File = open('automate_online-materials/meetingminutes.pdf', 'rb')
pdf2File = open('automate_online-materials/meetingminutes2.pdf', 'rb')
pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
pdf2Reader = PyPDF2.PdfFileReader(pdf2File)

# initiate pdfWriter
pdfWriter = PyPDF2.PdfFileWriter()

# for each file for each page we want to add the page to the pdfWriter
for page in range(pdf1Reader.numPages):
    pageObj = pdf1Reader.getPage(page)
    pdfWriter.addPage(pageObj)

for page in range(pdf2Reader.numPages):
    pageObj = pdf2Reader.getPage(page)
    pdfWriter.addPage(pageObj)

# note now it's WRITE binary, used to be read binary
pdfOutputFile = open('combinedminutes.pdf', 'wb')
# this thing can only add apges TO THE END - cant insert them in the middle
pdfWriter.write(pdfOutputFile)

pdfOutputFile.close()
pdf1File.close()
pdf2File.close()


print("-------------------- ROTATING PAGES --------------------")
# options: 90, 180, 270, etc

minutesFile = open('combinedminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
page = pdfReader.getPage(0)
page.rotateClockwise(90)

pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(page)
pdfOutputFile = open('rotated_minutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)

pdfOutputFile.close()
minutesFile.close()


print("-------------------- OVERLAYING --------------------")
minutesFile = open('combinedminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(minutesFile)
m_page = pdfReader.getPage(0)

watermarkFile = open('automate_online-materials/watermark.pdf', 'rb')
pdfWatermarkReader = PyPDF2.PdfFileReader(watermarkFile)
w_page = pdfWatermarkReader.getPage(0)

# KEY
m_page.mergePage(w_page)

pdfWriter = PyPDF2.PdfFileWriter()
pdfWriter.addPage(m_page)
pdfOutputFile = open('merged_minutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)

pdfOutputFile.close()
minutesFile.close()
watermarkFile.close()


print("-------------------- ENCRYPTING --------------------")
pdfFile = open('combinedminutes.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFile)
pdfWriter = PyPDF2.PdfFileWriter()

for p in range(pdfReader.numPages):
    pdfWriter.addPage(pdfReader.getPage(0))

# KEY
pdfWriter.encrypt('swordfish')
pdfOutputFile = open('encrypted_minutes.pdf', 'wb')
pdfWriter.write(pdfOutputFile)

pdfOutputFile.close()
pdfFile.close()


# ==============================================================================
# Now word docs!
print("-------------------- WORD --------------------")

# docx structure:
# highest level = Docubment object
# second down = Paragraph object (triggered when user hits return)
# third and final = Run object. A "run" is a continuous streak of similarly formatted text (same "style"). Eg if you had 4 words in normal, then 2 in bold, then 2 in italic, you'd have 3 runs of size 4,2,2.

doc = docx.Document('automate_online-materials/demo.docx')
print(len(doc.paragraphs))

print(doc.paragraphs[0].text)  # use text to get text
print(doc.paragraphs[1].text)
print(doc.paragraphs[1].runs[0].text)  # each para has runs
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)

# to get just the text from a word file (ignoring all the styling info) use getText()


print("-------------------- GET TEXT --------------------")


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for p in doc.paragraphs:
        # fullText.append(p.text)
        fullText.append('  ' + p.text)  # if we wanted to also indent each para!
    return '\n'.join(fullText)


print(getText('automate_online-materials/demo.docx'))


print("-------------------- STYLE --------------------")
# there are 3 types of styles in word:
# 1 paragraph styles = applied to paras
# 2 character styles = applied to runs
# 3 linked styles = applied to both
# you can give both Paragraph and Run objects styles by setting their style attribute to a string
# When using a linked style for a Run object, you will need to add ' Char' to the end of its name. For example, to set the Quote linked style for a Paragraph object, you would use paragraphObj.style = 'Quote', but for a Run object, you would use runObj.style = 'Quote Char'.
# NOTE: in current version of python-docx the styles used are default styles / those saved with the .docx. New can't be added.

# run attributes:
# run styles can further have run attributes (eg bold, italic, all_caps, etc)
# attributes can have one of 3 values attached to them
# True = by default on
# Flase = by default off
# None = defaults to whatever the run's style is set to

print(doc.paragraphs[0].text)
doc.paragraphs[0].style = 'Heading 1'
print(doc.paragraphs[1].text)
# if we didn't add "Char" at the end we'd get an error
doc.paragraphs[1].runs[0].style = 'Heading 2 Char'
doc.paragraphs[1].runs[1].style = 'Intense Quote Char'
doc.paragraphs[1].runs[1].underline = True

doc.save('restyled.docx')


print("-------------------- WRITING WORD --------------------")

doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj2 = doc.add_paragraph('This is a second paragraph!')
paraObj2.add_run('this is text being added to second parapgrah!')

# both add_para and add_run accept a second optional argument that is a string telling the style
doc.add_paragraph('This is gonna be loud.', 'Heading 1')

# we can add headings too
# NOTE headings also count as paragraphs, when counting which you're on
doc.add_heading('Jumbaaaa ymbaaa', 2)

paraObj3 = doc.add_paragraph('Need more text.')
paraObj3.add_run('Some more!')
paraObj3.add_run('And more!')

# a line break after 3rd para 1st run
doc.paragraphs[4].runs[0].add_break()

# and a page break after 2nd run
doc.paragraphs[4].runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)

# add a picture
doc.add_picture('sourdough_starter.png', width=docx.shared.Inches(1))

doc.save('greetings_humanity.docx')
